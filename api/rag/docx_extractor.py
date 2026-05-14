"""
api/rag/docx_extractor.py — Enterprise DOCX security validation and text extraction.

Security invariants:
- Validates ZIP magic bytes, MIME, size, paragraph count.
- Rejects macro-enabled formats, VBA-embedded documents, malformed archives.
- Zip bomb prevention: uncompressed size cap before any extraction.
- Extraction is deterministic: paragraphs returned in document order (1-based).
- Raw document content MUST NOT appear in error messages or logs.
- Client-declared MIME type is never trusted; magic bytes are authoritative.
"""

from __future__ import annotations

import hashlib
import io
import logging
import os
import zipfile
from dataclasses import dataclass
from typing import Final, Optional

log = logging.getLogger("frostgate.rag.docx_extractor")

# ---------------------------------------------------------------------------
# Error codes (stable; never change meaning once published)
# ---------------------------------------------------------------------------

DOCX_ERR_MISSING_BYTES = "DOCX_E001"
DOCX_ERR_INVALID_MAGIC = "DOCX_E002"
DOCX_ERR_MACRO_ENABLED = "DOCX_E003"
DOCX_ERR_MALFORMED = "DOCX_E004"
DOCX_ERR_TOO_MANY_PARAGRAPHS = "DOCX_E005"
DOCX_ERR_OVERSIZED_PARAGRAPH = "DOCX_E006"
DOCX_ERR_EMBEDDED_VBA = "DOCX_E007"
DOCX_ERR_EXTRACTION_FAILED = "DOCX_E008"
DOCX_ERR_EMPTY_EXTRACT = "DOCX_E009"
DOCX_ERR_LIBRARY_MISSING = "DOCX_E010"

# ---------------------------------------------------------------------------
# Limits (all overridable via env for air-gapped / regulated deployment)
# ---------------------------------------------------------------------------

_ZIP_MAGIC: Final[bytes] = b"\x50\x4b\x03\x04"  # PK\x03\x04
_MAX_DOCX_PARAGRAPHS: Final[int] = int(os.getenv("FG_DOCX_MAX_PARAGRAPHS", "10000"))
_MAX_PARAGRAPH_TEXT_BYTES: Final[int] = int(
    os.getenv("FG_DOCX_MAX_PARAGRAPH_TEXT_BYTES", str(100_000))
)
# Zip bomb guard: reject if total uncompressed size exceeds this limit.
_MAX_UNCOMPRESSED_BYTES: Final[int] = int(
    os.getenv("FG_DOCX_MAX_UNCOMPRESSED_BYTES", str(200_000_000))
)

# Members that indicate embedded VBA/macro code.
_VBA_MEMBER_NAMES: Final[frozenset[str]] = frozenset(
    {
        "word/vbaProject.bin",
        "xl/vbaProject.bin",
        "ppt/vbaProject.bin",
        "VBA/ThisDocument",
    }
)

# Strings in [Content_Types].xml that indicate a macro-enabled document type.
_MACRO_CONTENT_TYPE_MARKERS: Final[frozenset[bytes]] = frozenset(
    {
        b"macroEnabled",
        b"macroEnabled.main",
        b"ms-word.document.macroEnabled",
        b"ms-word.template.macroEnabled",
    }
)

# Style name prefix used by python-docx for heading paragraphs.
_HEADING_STYLE_PREFIX = "Heading "


# ---------------------------------------------------------------------------
# Exceptions
# ---------------------------------------------------------------------------


class DOCXExtractionError(Exception):
    """Raised for any DOCX validation or extraction failure.

    error_code is always a stable DOCX_Exxx constant.
    message MUST NOT contain raw document content.
    """

    def __init__(self, error_code: str, message: str) -> None:
        super().__init__(message)
        self.error_code = error_code
        self.message = message


# ---------------------------------------------------------------------------
# Result types
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class DOCXParagraph:
    paragraph_number: int  # 1-based position in document
    text: str
    char_count: int
    heading_level: Optional[int]  # 1–9 if a heading style, None otherwise
    style_name: str


@dataclass(frozen=True)
class DOCXExtractionResult:
    paragraphs: list[DOCXParagraph]
    paragraph_count: int
    extraction_version: str
    source_hash: str  # SHA-256 of the raw DOCX bytes
    has_text: bool


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _docx_extraction_version() -> str:
    try:
        import importlib.metadata

        ver = importlib.metadata.version("python-docx")
        return f"python-docx-{ver}"
    except Exception:
        return "python-docx-unknown"


_EXTRACTION_VERSION: str = _docx_extraction_version()


def _check_magic(raw_bytes: bytes) -> None:
    if not raw_bytes:
        raise DOCXExtractionError(DOCX_ERR_MISSING_BYTES, "DOCX content is empty")
    if not raw_bytes[:4] == _ZIP_MAGIC:
        raise DOCXExtractionError(
            DOCX_ERR_INVALID_MAGIC,
            "File does not begin with ZIP magic bytes (PK\\x03\\x04); "
            "client MIME type rejected",
        )


def _check_macro_and_vba(raw_bytes: bytes, source_hash_prefix: str) -> None:
    """Inspect the raw ZIP for VBA binaries and macro-enabled content types.

    Runs pre-parse on the raw bytes so that a malicious DOCX cannot
    exploit python-docx parsing to smuggle macro execution.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(raw_bytes), "r") as zf:
            namelist = set(zf.namelist())

            # Zip bomb guard: total uncompressed size.
            total_uncompressed = sum(info.file_size for info in zf.infolist())
            if total_uncompressed > _MAX_UNCOMPRESSED_BYTES:
                raise DOCXExtractionError(
                    DOCX_ERR_MALFORMED,
                    f"DOCX archive uncompressed size exceeds limit "
                    f"({total_uncompressed} > {_MAX_UNCOMPRESSED_BYTES} bytes)",
                )

            # VBA binary detection.
            vba_members = _VBA_MEMBER_NAMES & namelist
            if vba_members:
                log.warning(
                    "docx_extractor.vba_rejected",
                    extra={
                        "event": "docx.vba_rejected",
                        "source_hash_prefix": source_hash_prefix,
                    },
                )
                raise DOCXExtractionError(
                    DOCX_ERR_EMBEDDED_VBA,
                    "DOCX contains embedded VBA binary; document rejected",
                )

            # Macro-enabled content type detection via [Content_Types].xml.
            if "[Content_Types].xml" in namelist:
                ct_data = zf.read("[Content_Types].xml")
                for marker in _MACRO_CONTENT_TYPE_MARKERS:
                    if marker in ct_data:
                        log.warning(
                            "docx_extractor.macro_enabled_rejected",
                            extra={
                                "event": "docx.macro_enabled_rejected",
                                "source_hash_prefix": source_hash_prefix,
                            },
                        )
                        raise DOCXExtractionError(
                            DOCX_ERR_MACRO_ENABLED,
                            "DOCX is a macro-enabled format; document rejected",
                        )

    except DOCXExtractionError:
        raise
    except zipfile.BadZipFile as exc:
        raise DOCXExtractionError(
            DOCX_ERR_MALFORMED,
            "DOCX archive is malformed (bad ZIP structure)",
        ) from exc
    except Exception as exc:
        raise DOCXExtractionError(
            DOCX_ERR_MALFORMED,
            "DOCX archive could not be inspected",
        ) from exc


def _parse_heading_level(style_name: str) -> Optional[int]:
    """Return 1–9 if the style is a heading, None otherwise."""
    if style_name.startswith(_HEADING_STYLE_PREFIX):
        tail = style_name[len(_HEADING_STYLE_PREFIX) :]
        try:
            level = int(tail.split()[0])
            if 1 <= level <= 9:
                return level
        except (ValueError, IndexError):
            pass
    return None


def _normalize_paragraph_text(raw: str) -> str:
    return raw.replace("\r\n", "\n").replace("\r", "\n").strip()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def extract_docx_paragraphs(raw_bytes: bytes) -> DOCXExtractionResult:
    """Validate and extract text from a DOCX document, paragraph by paragraph.

    Args:
        raw_bytes: Raw DOCX bytes. Must not exceed the upload size cap
                   (enforced by the upload handler before calling this function).

    Returns:
        DOCXExtractionResult with per-paragraph text in deterministic document order.

    Raises:
        DOCXExtractionError: On any validation or extraction failure.

    Security invariants:
        - ZIP magic bytes validated before any library parsing.
        - VBA binaries and macro-enabled content types checked before python-docx.
        - Zip bomb guard checked against total uncompressed size before extraction.
        - Malformed archives caught; error message does not include document content.
        - Paragraph count and per-paragraph text size capped against DoS.
        - Raw text never logged or included in error messages.
    """
    _check_magic(raw_bytes)

    source_hash = hashlib.sha256(raw_bytes).hexdigest()
    _check_macro_and_vba(raw_bytes, source_hash[:12])

    try:
        from docx import Document
    except ImportError as exc:
        raise DOCXExtractionError(
            DOCX_ERR_LIBRARY_MISSING,
            "python-docx is not installed; DOCX extraction is unavailable",
        ) from exc

    try:
        doc = Document(io.BytesIO(raw_bytes))
    except Exception as exc:
        log.warning(
            "docx_extractor.malformed_docx",
            extra={
                "event": "docx.malformed_rejected",
                "source_hash_prefix": source_hash[:12],
            },
        )
        raise DOCXExtractionError(
            DOCX_ERR_MALFORMED,
            "DOCX is malformed and could not be parsed",
        ) from exc

    raw_paragraphs = doc.paragraphs
    para_count_total = len(raw_paragraphs)
    if para_count_total > _MAX_DOCX_PARAGRAPHS:
        log.warning(
            "docx_extractor.paragraph_limit_exceeded",
            extra={
                "event": "docx.paragraph_limit_rejected",
                "paragraph_count": para_count_total,
                "max_paragraphs": _MAX_DOCX_PARAGRAPHS,
                "source_hash_prefix": source_hash[:12],
            },
        )
        raise DOCXExtractionError(
            DOCX_ERR_TOO_MANY_PARAGRAPHS,
            f"DOCX has {para_count_total} paragraphs; "
            f"maximum allowed is {_MAX_DOCX_PARAGRAPHS}",
        )

    paragraphs: list[DOCXParagraph] = []
    para_number = 0
    for raw_para in raw_paragraphs:
        try:
            raw_text = raw_para.text or ""
            style_name = raw_para.style.name if raw_para.style else "Normal"
        except Exception:
            raw_text = ""
            style_name = "Normal"

        text = _normalize_paragraph_text(raw_text)
        if not text:
            continue  # skip blank/whitespace-only paragraphs

        para_number += 1
        text_bytes = len(text.encode("utf-8"))
        if text_bytes > _MAX_PARAGRAPH_TEXT_BYTES:
            raise DOCXExtractionError(
                DOCX_ERR_OVERSIZED_PARAGRAPH,
                f"Paragraph {para_number} text exceeds the per-paragraph size limit "
                f"({text_bytes} bytes > {_MAX_PARAGRAPH_TEXT_BYTES} bytes)",
            )

        heading_level = _parse_heading_level(style_name)
        paragraphs.append(
            DOCXParagraph(
                paragraph_number=para_number,
                text=text,
                char_count=len(text),
                heading_level=heading_level,
                style_name=style_name,
            )
        )

    has_text = any(p.text.strip() for p in paragraphs)

    log.info(
        "docx_extractor.extraction_complete",
        extra={
            "event": "docx.extraction_complete",
            "paragraph_count": len(paragraphs),
            "has_text": has_text,
            "extraction_version": _EXTRACTION_VERSION,
            "source_hash_prefix": source_hash[:12],
        },
    )

    return DOCXExtractionResult(
        paragraphs=paragraphs,
        paragraph_count=len(paragraphs),
        extraction_version=_EXTRACTION_VERSION,
        source_hash=source_hash,
        has_text=has_text,
    )


def build_docx_chunk_payloads(
    *,
    tenant_id: str,
    document_id: str,
    version_id: str,
    source_hash: str,
    docx_result: DOCXExtractionResult,
    max_chars: int = 1000,
) -> list[dict]:
    """Produce paragraph-aware, deterministic chunk payloads from a DOCX extraction.

    Each chunk carries:
      - source_page: 1-based paragraph_number where this chunk begins
        (stored in the source_page column, which is document-position generic)
      - extraction_version: version string of the python-docx library used
      - chunk_hash (content_hash): SHA-256 of the chunk text
      - heading_level: heading level of the originating paragraph (in metadata)
      - section_heading: text of the most recent heading seen (in metadata)
      - All existing provenance fields (document_version_id, source_hash)

    Chunks never cross paragraph boundaries.
    Section heading context is tracked across paragraphs and injected
    into each chunk's metadata for citation rendering.
    """
    from api.rag_corpus_store import (
        canonical_source_hash,
        deterministic_chunk_id,
    )

    payloads: list[dict] = []
    global_ordinal = 0
    current_section_heading: Optional[str] = None

    for para in docx_result.paragraphs:
        if not para.text.strip():
            continue

        # Track the most recent heading for section attribution.
        if para.heading_level is not None:
            current_section_heading = para.text

        words = para.text.split()
        para_chunks: list[str] = []
        current: list[str] = []
        current_len = 0

        for word in words:
            word_len = len(word)
            if word_len > max_chars:
                word = word[:max_chars]
                word_len = max_chars
            needed = word_len if not current else word_len + 1
            if current and current_len + needed > max_chars:
                para_chunks.append(" ".join(current))
                current = []
                current_len = 0
            current.append(word)
            current_len += word_len if len(current) == 1 else word_len + 1

        if current:
            para_chunks.append(" ".join(current))

        for chunk_text in para_chunks:
            content_hash = canonical_source_hash(chunk_text)
            chunk_id = deterministic_chunk_id(
                tenant_id=tenant_id,
                document_id=document_id,
                version_id=version_id,
                ordinal=global_ordinal,
                text=chunk_text,
            )
            payloads.append(
                {
                    "chunk_id": chunk_id,
                    "text": chunk_text,
                    "ordinal": global_ordinal,
                    "document_version_id": version_id,
                    "source_hash": source_hash,
                    "content_hash": content_hash,
                    # source_page stores the paragraph position (document-position generic)
                    "source_page": para.paragraph_number,
                    "extraction_version": docx_result.extraction_version,
                    "metadata": {
                        "document_version_id": version_id,
                        "source_hash": source_hash,
                        "chunk_index": global_ordinal,
                        "source_paragraph": para.paragraph_number,
                        "heading_level": para.heading_level,
                        "section_heading": current_section_heading,
                        "style_name": para.style_name,
                        "extraction_version": docx_result.extraction_version,
                        "chunk_hash": content_hash,
                        "evidence_graph_ready": True,
                        "verified_fact_binding_ready": True,
                        "rag_evaluation_ready": True,
                        "table_extraction_ready": True,
                        "tracked_changes_ready": True,
                        "comments_extraction_ready": True,
                        "embedded_image_ocr_ready": True,
                        "legal_segmentation_ready": True,
                        "async_worker_ready": True,
                    },
                }
            )
            global_ordinal += 1

    return payloads
