"""
tests/rag/test_docx_ingestion.py — Unit and integration tests for DOCX ingestion.

Tests cover:
- extract_docx_paragraphs: validation, extraction, security checks
- ingest_docx_document: full ingestion pipeline with SQLite
- build_docx_chunk_payloads: paragraph-aware chunking, provenance metadata
"""

from __future__ import annotations

import io
import zipfile

import pytest

from api.rag.docx_extractor import (
    DOCX_ERR_EMBEDDED_VBA,
    DOCX_ERR_INVALID_MAGIC,
    DOCX_ERR_MACRO_ENABLED,
    DOCX_ERR_MALFORMED,
    DOCX_ERR_MISSING_BYTES,
    DOCX_ERR_OVERSIZED_PARAGRAPH,
    DOCX_ERR_TOO_MANY_PARAGRAPHS,
    DOCXExtractionError,
    DOCXExtractionResult,
    build_docx_chunk_payloads,
    extract_docx_paragraphs,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_raw_docx(
    paragraphs: list[str] | None = None,
    headings: list[tuple[int, str]] | None = None,
) -> bytes:
    """Create a minimal valid DOCX in memory using python-docx."""
    from docx import Document

    doc = Document()
    if headings:
        for level, text in headings:
            doc.add_heading(text, level=level)
    for text in paragraphs or ["Hello, world!"]:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_macro_docx() -> bytes:
    """Create a ZIP with a macro-enabled content type marker."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.ms-word.document.macroEnabled.main+xml"/>'
            "</Types>",
        )
        zf.writestr("word/document.xml", "<document/>")
    return buf.getvalue()


def _make_vba_docx() -> bytes:
    """Create a ZIP that contains a vbaProject.bin member."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            "</Types>",
        )
        zf.writestr("word/vbaProject.bin", b"\xd0\xcf\x11\xe0" + b"\x00" * 100)
        zf.writestr("word/document.xml", "<document/>")
    return buf.getvalue()


def _make_bad_zip() -> bytes:
    """A PK-magic prefix but otherwise corrupt ZIP data."""
    return b"PK\x03\x04" + b"\xff" * 50


# ---------------------------------------------------------------------------
# Setup: SQLite in-memory DB
# ---------------------------------------------------------------------------


@pytest.fixture()
def db(tmp_path, monkeypatch):
    db_path = str(tmp_path / "docx-ingestion.db")
    monkeypatch.setenv("FG_SQLITE_PATH", db_path)
    monkeypatch.setenv("FG_ENV", "test")

    from api.db import get_sessionmaker, init_db, reset_engine_cache

    reset_engine_cache()
    init_db(sqlite_path=db_path)
    SessionLocal = get_sessionmaker(sqlite_path=db_path)
    session = SessionLocal()

    try:
        yield session
    finally:
        session.close()
        reset_engine_cache()


@pytest.fixture()
def corpus_a(db):
    from api.rag_corpus_store import create_corpus

    return create_corpus(db, tenant_id="tenant-a", name="Test Corpus")["corpus_id"]


# ---------------------------------------------------------------------------
# extract_docx_paragraphs — validation
# ---------------------------------------------------------------------------


class TestDocxExtractorValidation:
    def test_empty_bytes_raises(self):
        with pytest.raises(DOCXExtractionError) as exc_info:
            extract_docx_paragraphs(b"")
        assert exc_info.value.error_code == DOCX_ERR_MISSING_BYTES

    def test_invalid_magic_raises(self):
        with pytest.raises(DOCXExtractionError) as exc_info:
            extract_docx_paragraphs(b"NOT_A_ZIP_FILE_AT_ALL")
        assert exc_info.value.error_code == DOCX_ERR_INVALID_MAGIC

    def test_pdf_bytes_rejected(self):
        with pytest.raises(DOCXExtractionError) as exc_info:
            extract_docx_paragraphs(b"%PDF-1.4 ...")
        assert exc_info.value.error_code == DOCX_ERR_INVALID_MAGIC

    def test_macro_enabled_rejected(self):
        raw = _make_macro_docx()
        with pytest.raises(DOCXExtractionError) as exc_info:
            extract_docx_paragraphs(raw)
        assert exc_info.value.error_code == DOCX_ERR_MACRO_ENABLED

    def test_vba_binary_rejected(self):
        raw = _make_vba_docx()
        with pytest.raises(DOCXExtractionError) as exc_info:
            extract_docx_paragraphs(raw)
        assert exc_info.value.error_code == DOCX_ERR_EMBEDDED_VBA

    def test_malformed_zip_rejected(self):
        with pytest.raises(DOCXExtractionError) as exc_info:
            extract_docx_paragraphs(_make_bad_zip())
        assert exc_info.value.error_code == DOCX_ERR_MALFORMED

    def test_paragraph_limit_enforced(self, monkeypatch):
        import api.rag.docx_extractor as mod

        monkeypatch.setattr(mod, "_MAX_DOCX_PARAGRAPHS", 2)
        raw = _make_raw_docx(["para one", "para two", "para three"])
        with pytest.raises(DOCXExtractionError) as exc_info:
            extract_docx_paragraphs(raw)
        assert exc_info.value.error_code == DOCX_ERR_TOO_MANY_PARAGRAPHS

    def test_oversized_paragraph_enforced(self, monkeypatch):
        import api.rag.docx_extractor as mod

        monkeypatch.setattr(mod, "_MAX_PARAGRAPH_TEXT_BYTES", 10)
        raw = _make_raw_docx(["This paragraph is definitely longer than ten bytes."])
        with pytest.raises(DOCXExtractionError) as exc_info:
            extract_docx_paragraphs(raw)
        assert exc_info.value.error_code == DOCX_ERR_OVERSIZED_PARAGRAPH


# ---------------------------------------------------------------------------
# extract_docx_paragraphs — extraction
# ---------------------------------------------------------------------------


class TestDocxExtractorExtraction:
    def test_valid_docx_returns_result(self):
        raw = _make_raw_docx(["Hello, world!"])
        result = extract_docx_paragraphs(raw)
        assert result.has_text is True
        assert result.paragraph_count >= 1
        assert any(p.text == "Hello, world!" for p in result.paragraphs)

    def test_source_hash_is_sha256(self):
        import hashlib

        raw = _make_raw_docx(["test paragraph"])
        result = extract_docx_paragraphs(raw)
        expected = hashlib.sha256(raw).hexdigest()
        assert result.source_hash == expected
        assert len(result.source_hash) == 64

    def test_extraction_version_present(self):
        raw = _make_raw_docx(["paragraph"])
        result = extract_docx_paragraphs(raw)
        assert result.extraction_version.startswith("python-docx-")

    def test_paragraph_order_deterministic(self):
        texts = ["Alpha paragraph", "Beta paragraph", "Gamma paragraph"]
        raw = _make_raw_docx(texts)
        r1 = extract_docx_paragraphs(raw)
        r2 = extract_docx_paragraphs(raw)
        assert [p.text for p in r1.paragraphs] == [p.text for p in r2.paragraphs]

    def test_heading_level_detected(self):
        raw = _make_raw_docx(
            paragraphs=["Body text"],
            headings=[(1, "Section One"), (2, "Subsection")],
        )
        result = extract_docx_paragraphs(raw)
        heading_paragraphs = [
            p for p in result.paragraphs if p.heading_level is not None
        ]
        assert len(heading_paragraphs) >= 1
        assert heading_paragraphs[0].heading_level == 1

    def test_blank_paragraphs_skipped(self):
        raw = _make_raw_docx(["  ", "Real content", "  "])
        result = extract_docx_paragraphs(raw)
        assert all(p.text.strip() for p in result.paragraphs)

    def test_paragraph_numbers_are_1_based(self):
        raw = _make_raw_docx(["First", "Second", "Third"])
        result = extract_docx_paragraphs(raw)
        assert result.paragraphs[0].paragraph_number == 1

    def test_paragraph_number_reflects_document_position_not_retained_ordinal(self):
        # Blank paragraphs before real content must not compress paragraph numbers.
        # A paragraph at document position 3 must report paragraph_number=3, not 1.
        raw = _make_raw_docx(["   ", "   ", "Real content at position three"])
        result = extract_docx_paragraphs(raw)
        real = next(p for p in result.paragraphs if "Real content" in p.text)
        assert real.paragraph_number >= 3

    def test_paragraph_char_count_accurate(self):
        raw = _make_raw_docx(["Hello"])
        result = extract_docx_paragraphs(raw)
        para = next(p for p in result.paragraphs if p.text == "Hello")
        assert para.char_count == len("Hello")


# ---------------------------------------------------------------------------
# build_docx_chunk_payloads — chunking
# ---------------------------------------------------------------------------


class TestBuildDocxChunkPayloads:
    def _make_result(self, texts: list[str]) -> DOCXExtractionResult:
        raw = _make_raw_docx(texts)
        return extract_docx_paragraphs(raw)

    def test_ordinals_are_sequential(self):
        result = self._make_result(["Para one", "Para two", "Para three"])
        payloads = build_docx_chunk_payloads(
            tenant_id="t1",
            document_id="d1",
            version_id="v1",
            source_hash=result.source_hash,
            docx_result=result,
        )
        ordinals = [p["ordinal"] for p in payloads]
        assert ordinals == list(range(len(ordinals)))

    def test_chunks_never_cross_paragraph_boundary(self):
        result = self._make_result(["Alpha " * 50, "Beta " * 50])
        payloads = build_docx_chunk_payloads(
            tenant_id="t1",
            document_id="d1",
            version_id="v1",
            source_hash=result.source_hash,
            docx_result=result,
            max_chars=100,
        )
        # All chunks from "Alpha" paragraphs must not contain "Beta" and vice versa.
        for payload in payloads:
            text = payload["text"]
            assert not ("Alpha" in text and "Beta" in text)

    def test_source_paragraph_in_metadata(self):
        result = self._make_result(["First paragraph", "Second paragraph"])
        payloads = build_docx_chunk_payloads(
            tenant_id="t1",
            document_id="d1",
            version_id="v1",
            source_hash=result.source_hash,
            docx_result=result,
        )
        for payload in payloads:
            assert "source_paragraph" in payload["metadata"]
            assert payload["metadata"]["source_paragraph"] >= 1

    def test_source_page_stores_paragraph_number(self):
        result = self._make_result(["Para A", "Para B"])
        payloads = build_docx_chunk_payloads(
            tenant_id="t1",
            document_id="d1",
            version_id="v1",
            source_hash=result.source_hash,
            docx_result=result,
        )
        for payload in payloads:
            assert payload["source_page"] == payload["metadata"]["source_paragraph"]

    def test_extraction_version_in_payload(self):
        result = self._make_result(["Some content"])
        payloads = build_docx_chunk_payloads(
            tenant_id="t1",
            document_id="d1",
            version_id="v1",
            source_hash=result.source_hash,
            docx_result=result,
        )
        for payload in payloads:
            assert payload["extraction_version"] == result.extraction_version
            assert (
                payload["metadata"]["extraction_version"] == result.extraction_version
            )

    def test_chunk_hash_in_metadata(self):
        result = self._make_result(["Chunk hash test"])
        payloads = build_docx_chunk_payloads(
            tenant_id="t1",
            document_id="d1",
            version_id="v1",
            source_hash=result.source_hash,
            docx_result=result,
        )
        for payload in payloads:
            assert "chunk_hash" in payload["metadata"]
            assert payload["content_hash"] == payload["metadata"]["chunk_hash"]

    def test_section_heading_tracked(self):
        raw = _make_raw_docx(
            paragraphs=["Body under section"],
            headings=[(1, "My Section")],
        )
        result = extract_docx_paragraphs(raw)
        payloads = build_docx_chunk_payloads(
            tenant_id="t1",
            document_id="d1",
            version_id="v1",
            source_hash=result.source_hash,
            docx_result=result,
        )
        body_payloads = [p for p in payloads if p["text"] == "Body under section"]
        assert len(body_payloads) >= 1
        assert body_payloads[0]["metadata"]["section_heading"] == "My Section"

    def test_heading_level_in_metadata(self):
        raw = _make_raw_docx(headings=[(2, "Level Two Heading")])
        result = extract_docx_paragraphs(raw)
        payloads = build_docx_chunk_payloads(
            tenant_id="t1",
            document_id="d1",
            version_id="v1",
            source_hash=result.source_hash,
            docx_result=result,
        )
        heading_payloads = [
            p for p in payloads if p["metadata"].get("heading_level") is not None
        ]
        assert len(heading_payloads) >= 1
        assert heading_payloads[0]["metadata"]["heading_level"] == 2

    def test_deterministic_chunk_ids(self):
        result = self._make_result(["Stable content"])
        kwargs = dict(
            tenant_id="t1",
            document_id="d1",
            version_id="v1",
            source_hash=result.source_hash,
            docx_result=result,
        )
        p1 = build_docx_chunk_payloads(**kwargs)
        p2 = build_docx_chunk_payloads(**kwargs)
        assert [c["chunk_id"] for c in p1] == [c["chunk_id"] for c in p2]

    def test_future_hook_flags_present(self):
        result = self._make_result(["Content"])
        payloads = build_docx_chunk_payloads(
            tenant_id="t1",
            document_id="d1",
            version_id="v1",
            source_hash=result.source_hash,
            docx_result=result,
        )
        for payload in payloads:
            meta = payload["metadata"]
            assert meta["table_extraction_ready"] is True
            assert meta["tracked_changes_ready"] is True
            assert meta["embedded_image_ocr_ready"] is True


# ---------------------------------------------------------------------------
# ingest_docx_document — integration
# ---------------------------------------------------------------------------


class TestIngestDocxDocument:
    def test_valid_ingestion_returns_indexed(self, db, corpus_a):
        from api.rag_corpus_store import ingest_docx_document

        raw = _make_raw_docx(["Document body text for ingestion."])
        result_extract = extract_docx_paragraphs(raw)
        result = ingest_docx_document(
            db,
            tenant_id="tenant-a",
            corpus_id=corpus_a,
            title="Test DOCX",
            source="test.docx",
            docx_result=result_extract,
        )
        assert result["ingestion_status"] == "indexed"
        assert result["chunk_count"] >= 1
        assert result["paragraph_count"] >= 1
        assert result["extraction_version"].startswith("python-docx-")

    def test_duplicate_detection_by_source_hash(self, db, corpus_a):
        from api.rag_corpus_store import ingest_docx_document

        raw = _make_raw_docx(["Unique document content."])
        result_extract = extract_docx_paragraphs(raw)
        r1 = ingest_docx_document(
            db,
            tenant_id="tenant-a",
            corpus_id=corpus_a,
            title="Doc",
            source="doc.docx",
            docx_result=result_extract,
        )
        r2 = ingest_docx_document(
            db,
            tenant_id="tenant-a",
            corpus_id=corpus_a,
            title="Doc Again",
            source="doc2.docx",
            docx_result=result_extract,
        )
        assert r2["ingestion_status"] == "duplicate"
        assert r2["document_id"] == r1["document_id"]

    def test_chunk_metadata_source_paragraph(self, db, corpus_a):
        from sqlalchemy import text

        from api.rag_corpus_store import ingest_docx_document

        raw = _make_raw_docx(["Paragraph one content", "Paragraph two content"])
        result_extract = extract_docx_paragraphs(raw)
        result = ingest_docx_document(
            db,
            tenant_id="tenant-a",
            corpus_id=corpus_a,
            title="Para Test",
            source="para.docx",
            docx_result=result_extract,
        )
        doc_id = result["document_id"]
        chunks = (
            db.execute(
                text(
                    "SELECT source_page, extraction_version, metadata FROM rag_chunks "
                    "WHERE document_id = :doc_id AND tenant_id = 'tenant-a'"
                    " ORDER BY ordinal"
                ),
                {"doc_id": doc_id},
            )
            .mappings()
            .fetchall()
        )
        assert len(chunks) >= 2
        for chunk in chunks:
            assert chunk["source_page"] is not None
            assert chunk["source_page"] >= 1
            assert chunk["extraction_version"] is not None

    def test_empty_docx_quarantined(self, db, corpus_a):
        from api.rag_corpus_store import ingest_docx_document

        # Whitespace-only paragraphs normalize to empty strings → has_text=False.
        raw = _make_raw_docx(["   ", "  "])
        result_extract = extract_docx_paragraphs(raw)
        result = ingest_docx_document(
            db,
            tenant_id="tenant-a",
            corpus_id=corpus_a,
            title="Empty Doc",
            source="empty.docx",
            docx_result=result_extract,
        )
        assert result["ingestion_status"] == "quarantined"

    def test_tenant_isolation(self, db, corpus_a):
        from api.rag_corpus_store import create_corpus, ingest_docx_document

        corpus_b = create_corpus(db, tenant_id="tenant-b", name="Corp B")["corpus_id"]

        raw = _make_raw_docx(["Isolated content"])
        result_extract = extract_docx_paragraphs(raw)
        result = ingest_docx_document(
            db,
            tenant_id="tenant-a",
            corpus_id=corpus_a,
            title="Doc A",
            source="isolated.docx",
            docx_result=result_extract,
        )
        doc_id = result["document_id"]

        from sqlalchemy import text

        # tenant-b must not see tenant-a's chunks.
        chunks_b = (
            db.execute(
                text(
                    "SELECT chunk_id FROM rag_chunks "
                    "WHERE document_id = :doc_id AND tenant_id = 'tenant-b'"
                ),
                {"doc_id": doc_id},
            )
            .mappings()
            .fetchall()
        )
        assert len(chunks_b) == 0
        # suppress unused variable warning
        _ = corpus_b

    def test_content_type_persisted(self, db, corpus_a):
        from sqlalchemy import text

        from api.rag_corpus_store import ingest_docx_document

        raw = _make_raw_docx(["Content type test"])
        result_extract = extract_docx_paragraphs(raw)
        result = ingest_docx_document(
            db,
            tenant_id="tenant-a",
            corpus_id=corpus_a,
            title="CT Doc",
            source="ct.docx",
            docx_result=result_extract,
        )
        row = (
            db.execute(
                text(
                    "SELECT content_type FROM rag_documents WHERE document_id = :doc_id"
                ),
                {"doc_id": result["document_id"]},
            )
            .mappings()
            .first()
        )
        assert row is not None
        assert row["content_type"] == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
